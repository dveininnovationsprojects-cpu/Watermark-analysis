from docx import Document
import pymupdf as fitz
import re
from io import BytesIO
import nltk 
from sentence_transformers import SentenceTransformer,util
import pytesseract
from PIL import Image

pytesseract.pytesseract.tesseract_cmd=r'C:/Program Files/Tesseract-OCR/tesseract.exe'
def detect_watermark_docx(doc,word):
    dic={} 
    doc=doc.read()
    doc=Document(BytesIO(doc))
    model = SentenceTransformer('all-MiniLM-L6-v2')
    text=''
    text2=''
    print(text)
    for para in doc.paragraphs:
            text_s = para.text.lower()
            for i in text_s:
                text+=i

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            text_s = para.text.lower()
            for i in text_s:
                text2+=i

    sentences = nltk.sent_tokenize(text)
    sentences2=nltk.sent_tokenize(text2)
    try:
        keywords=word.split(',')
    except:
        keywords=[word]
    
    sentence_embeddings = model.encode(sentences, convert_to_tensor=True)
    sentence_embeddings2=None
    if sentences2:
        sentence_embeddings2=model.encode(sentences2,convert_to_tensor=True)

    for kw in keywords:
        kw_emb = model.encode(kw, convert_to_tensor=True)
        scores = util.cos_sim(kw_emb, sentence_embeddings)[0]

        best_idx = scores.argmax()
        if scores[best_idx].item() > 0.5:
            dic['Keyword']=kw
            dic["Best match sentence"]= sentences[best_idx]
            dic["score"]=scores[best_idx].item()
    pattern = r"^(Mr|Mrs|Ms|Dr)\.\s([A-Z]\.)\s([A-Z]+)\s([A-Z]+)$"
    for kw in keywords:
        kw_emb = model.encode(kw, convert_to_tensor=True)
        if sentence_embeddings2:
            scores = util.cos_sim(kw_emb, sentence_embeddings2)[0]

        best_idx = scores.argmax()
        if scores[best_idx].item() > 0.5:
            dic['Keyword']=kw
            dic["Best match sentence"]= sentences[best_idx]
            dic["score"]=scores[best_idx].item()
                  
    dic2={}
    for section in doc.sections:
        header = section.header
        for rel in header.part.rels.values():
            if "image" in rel.reltype:
                dic2["Watermark image found:"]= rel.target_ref

    return dic,dic2

def detect_watermark_pdf(pdf,word):
    dic3={}
    pdf1=BytesIO(pdf.read())
    pdf1 = fitz.open(stream=pdf1.read(), filetype="pdf")
    for page in pdf1:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "opacity" in b and b["opacity"] < 1:
                dic3['transparent images']=True
    if not dic3:
        dic3['transparent images']=False
    

    model = SentenceTransformer('all-MiniLM-L6-v2')
    dic1={}
    pdf.seek(0)
    pdf_bytes = BytesIO(pdf.read())
    pdf2=fitz.open(stream=pdf_bytes,filetype='pdf')
    text=''
   
    for page in pdf2:
        tex=page.get_text()
        text += tex.replace("\x00", "").replace("\n", " ").strip()
        
    model = SentenceTransformer('all-MiniLM-L6-v2')
    sentences = nltk.sent_tokenize(text)
    try:
        keywords = word.split(',')
    except:
        keywords=[word.lower(),word.upper()]
    if sentences:
        sentence_embeddings = model.encode(sentences, convert_to_tensor=True)
    else:
        sentence_embeddings = None
    m=0
    if sentence_embeddings is not None:
        for kw in keywords:
            m+=1
            kw_emb = model.encode(kw, convert_to_tensor=True)
            scores = util.cos_sim(kw_emb, sentence_embeddings)[0]
            best_idx = scores.argmax()
            pattern = r"{}\s+((?:\S+\s+){{4}}\S+)".format(kw)
            print(text[:200])
            print(pattern)
            print(kw)
            print(re.escape(kw))
            if re.search(pattern, text):
                
                match=re.search(pattern,text)
                print(match)
                dic1['keyword'+str(m*10)]=match.group(0)
            dic1['Keyword'+str(m)]=kw
            dic1["Best match sentence"+str(m)]= sentences[best_idx]
            dic1["score"+str(m)]=scores[best_idx].item()

    dic2={}
    pdf.seek(0)  # important if coming from upload or buffer
    pdf3=BytesIO(pdf.read())
    doc = fitz.open(stream=pdf3, filetype="pdf")
    for page_index, page in enumerate(doc):
        text_dict = page.get_text("dict")
        for block in text_dict["blocks"]:
            if "lines" not in block:
                continue

            for line in block["lines"]:
                for span in line["spans"]:
                    if span["size"] > 40:   # font size check
                        dic2[page_index + 1] = span["text"]

    doc.close()


    return dic1,dic2,dic3

#print(detect_watermark_docx(open(r"C:\Users\lenovo\OneDrive\Documents\Indian Traditional knowledge.docx",'rb')))